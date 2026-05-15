"""
Script to apply Mayank's report formatting to the GuardNet report.
- Keeps all GuardNet text content UNCHANGED
- Applies: font, sizes, spacing, margins, heading styles, alignment, footer, colors
"""

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

INPUT_FILE = r'C:\Users\MohitTeli\Desktop\project-1\refrence_reports\GuardNet_Report (2).docx'
OUTPUT_FILE = r'C:\Users\MohitTeli\Desktop\project-1\refrence_reports\GuardNet_Report_Formatted.docx'

doc = Document(INPUT_FILE)

# ==========================================================
# 1. PAGE SETUP - Match Mayank's margins and page size
# ==========================================================
for section in doc.sections:
    section.page_width = Emu(7560310)    # Mayank's page width
    section.page_height = Emu(10692130)  # Mayank's page height
    section.top_margin = Emu(914400)     # ~0.72 inch
    section.bottom_margin = Emu(1371600) # ~1.08 inch
    section.left_margin = Emu(1141095)   # ~0.90 inch
    section.right_margin = Emu(683895)   # ~0.54 inch
    section.header_distance = Emu(450215)
    section.footer_distance = Emu(450215)

print("[OK] Page setup matched to Mayank's report.")

# ==========================================================
# 2. UPDATE DEFAULT STYLE (Normal) - Times New Roman, ~12pt, 1.1 line spacing
# ==========================================================
normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(12)
normal_style.font.color.rgb = RGBColor(0, 0, 0)  # Black text
normal_style.paragraph_format.line_spacing = 1.1
normal_style.paragraph_format.space_after = Pt(5)

# Force Times New Roman for East Asian fallback
rPr = normal_style.element.find(qn('w:rPr'))
if rPr is None:
    rPr = normal_style.element.makeelement(qn('w:rPr'), {})
    normal_style.element.append(rPr)
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')

print("[OK] Normal style set to Times New Roman 12pt, line spacing 1.1")

# ==========================================================
# 3. UPDATE HEADING STYLES - Match Mayank's formatting
# ==========================================================

# Heading 1: ~20pt, bold, black, center-ish (used for CHAPTER X titles)
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(20)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)  # Black, not blue
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(4)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Fix east-asian font
h1_rPr = h1.element.find(qn('w:rPr'))
if h1_rPr is not None:
    h1_rFonts = h1_rPr.find(qn('w:rFonts'))
    if h1_rFonts is not None:
        h1_rFonts.set(qn('w:ascii'), 'Times New Roman')
        h1_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        h1_rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        h1_rFonts.set(qn('w:cs'), 'Times New Roman')

# Heading 2: ~16pt, bold, black
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(16)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(8)
h2.paragraph_format.space_after = Pt(4)

h2_rPr = h2.element.find(qn('w:rPr'))
if h2_rPr is not None:
    h2_rFonts = h2_rPr.find(qn('w:rFonts'))
    if h2_rFonts is not None:
        h2_rFonts.set(qn('w:ascii'), 'Times New Roman')
        h2_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        h2_rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        h2_rFonts.set(qn('w:cs'), 'Times New Roman')

# Heading 3: ~14pt, bold, black
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(14)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after = Pt(4)

h3_rPr = h3.element.find(qn('w:rPr'))
if h3_rPr is not None:
    h3_rFonts = h3_rPr.find(qn('w:rFonts'))
    if h3_rFonts is not None:
        h3_rFonts.set(qn('w:ascii'), 'Times New Roman')
        h3_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        h3_rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        h3_rFonts.set(qn('w:cs'), 'Times New Roman')

print("[OK] Heading styles set to Times New Roman, bold, black (no blue).")

# ==========================================================
# 4. UPDATE List Paragraph STYLE
# ==========================================================
try:
    lp = doc.styles['List Paragraph']
    lp.font.name = 'Times New Roman'
    lp.font.size = Pt(12)
    lp.font.color.rgb = RGBColor(0, 0, 0)
    lp.paragraph_format.line_spacing = 1.1
    lp.paragraph_format.space_after = Pt(3)
    print("[OK] List Paragraph style updated.")
except:
    print("[SKIP] List Paragraph style not found.")

# ==========================================================
# 5. FORCE FONT ON EVERY PARAGRAPH AND RUN
#    (Override any run-level formatting that might ignore style)
# ==========================================================
for para in doc.paragraphs:
    # Set alignment for body text to JUSTIFY (like Mayank's report)
    style_name = para.style.name if para.style else ''
    
    if style_name in ('Normal', 'Normal (Web)', 'List Paragraph', 'Body Text'):
        # Only justify if not already center-aligned (preserving centered captions)
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            text_lower = para.text.strip().lower()
            # Don't justify figure/table captions that should stay centered
            if not (text_lower.startswith('fig') or text_lower.startswith('table')):
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Force Times New Roman on every run
    for run in para.runs:
        run.font.name = 'Times New Roman'
        # Set size only if not explicitly set for a reason (like cover page large text)
        if run.font.size is None:
            run.font.size = Pt(12)
        # Remove blue color from headings
        if style_name.startswith('Heading'):
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif run.font.color and run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
            # Keep black for body text
            if style_name not in ('Heading 1', 'Heading 2', 'Heading 3'):
                pass  # Leave any intentional coloring
        
        # Force east-asian font via XML
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Times New Roman')

print("[OK] All paragraphs and runs forced to Times New Roman, body text justified.")

# ==========================================================
# 6. UPDATE FOOTER TEXT
# ==========================================================
for section in doc.sections:
    footer = section.footer
    if footer:
        for para in footer.paragraphs:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

print("[OK] Footer formatting updated.")

# ==========================================================
# 7. FIX TABLE CELL FONTS
# ==========================================================
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    if run.font.size is None:
                        run.font.size = Pt(11)
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            rFonts.set(qn('w:ascii'), 'Times New Roman')
                            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            rFonts.set(qn('w:cs'), 'Times New Roman')

print("[OK] Table cell fonts forced to Times New Roman.")

# ==========================================================
# 8. SAVE OUTPUT
# ==========================================================
doc.save(OUTPUT_FILE)
print(f"\n=== DONE! Formatted report saved to: ===")
print(f"    {OUTPUT_FILE}")
print(f"\nContent is UNCHANGED. Only styling has been applied.")
